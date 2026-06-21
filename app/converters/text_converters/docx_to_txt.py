from pathlib import Path
from typing import Optional
from pydantic import BaseModel
from docx import Document
from app.converters.base import BaseConverter

class DocxToTxtConverter(BaseConverter):
    """Word (DOCX) to Plain Text."""

    @property
    def source_extension(self) -> str:
        return "docx"

    @property
    def target_extension(self) -> str:
        return "txt"

    def convert(self, input_path: Path, output_path: Path, options: Optional[BaseModel] = None) -> None:
        if not input_path.exists():
            raise FileNotFoundError(f"Input file not found: {input_path}")
        doc = Document(input_path)
        paragraphs_text = [p.text for p in doc.paragraphs]
        
        # Extract table values if tables exist
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                tables_text.append(row_text)
                
        full_text = "\n\n".join(p for p in paragraphs_text if p.strip())
        if tables_text:
            full_text += "\n\n=== Tables ===\n" + "\n".join(tables_text)
            
        output_path.write_text(full_text, encoding="utf-8")

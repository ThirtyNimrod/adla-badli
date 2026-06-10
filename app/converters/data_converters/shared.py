"""Shared helpers for the data converter group (csv, json sources).

All data conversions normalize input to a pandas DataFrame, then render
to the target format: openpyxl for styled XLSX, python-docx for Word
tables, and ReportLab platypus for paginated PDF tables.
"""
import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, letter, landscape
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle

# Warm-minimalism-adjacent table palette shared across renderers.
_HEADER_FILL_HEX = "F5F1EA"   # alabaster
_HEADER_TEXT_HEX = "2E2018"   # espresso
_BORDER_HEX = "D8CFC4"        # warm gray border

MAX_PREVIEW_COL_WIDTH = 48    # cap auto-fit column width in characters


def load_csv_dataframe(input_path: Path) -> pd.DataFrame:
    """Reads a CSV into a DataFrame, keeping all values intact."""
    df = pd.read_csv(input_path, dtype=object, keep_default_na=False)
    if df.empty and df.columns.empty:
        raise ValueError("CSV file contains no data.")
    return df


def load_json_dataframe(input_path: Path) -> pd.DataFrame:
    """Normalizes a JSON file into a flat DataFrame.

    Accepts: an array of objects (most common), a single object (one row),
    or an object whose first array-of-objects value is used as the table.
    Nested objects are flattened with dotted column names.
    """
    raw = json.loads(input_path.read_text(encoding="utf-8"))

    if isinstance(raw, list):
        data = raw
    elif isinstance(raw, dict):
        # Prefer the first list-of-dicts value as the tabular payload.
        list_values = [v for v in raw.values() if isinstance(v, list) and v and all(isinstance(i, dict) for i in v)]
        data = list_values[0] if list_values else [raw]
    else:
        raise ValueError("JSON root must be an object or an array of objects.")

    if not data:
        raise ValueError("JSON file contains no records.")

    df = pd.json_normalize(data)
    if df.empty:
        raise ValueError("JSON file produced no tabular data.")
    # Stringify residual complex cells (lists, nested dicts) for clean rendering.
    return df.map(lambda v: json.dumps(v, ensure_ascii=False) if isinstance(v, (list, dict)) else v)


def dataframe_to_xlsx(df: pd.DataFrame, output_path: Path, sheet_name: str = "Data") -> None:
    """Writes a DataFrame to XLSX with styled headers, frozen panes, auto-fit columns."""
    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        df.to_excel(writer, sheet_name=sheet_name, index=False)
        ws = writer.sheets[sheet_name]

        header_font = Font(bold=True, color=_HEADER_TEXT_HEX)
        header_fill = PatternFill(start_color=_HEADER_FILL_HEX, end_color=_HEADER_FILL_HEX, fill_type="solid")
        for cell in ws[1]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(vertical="center")

        ws.freeze_panes = "A2"

        for idx, column in enumerate(df.columns, start=1):
            longest = max([len(str(column))] + [len(str(v)) for v in df[column].head(200)])
            ws.column_dimensions[get_column_letter(idx)].width = min(longest + 3, MAX_PREVIEW_COL_WIDTH)


def dataframe_to_docx(df: pd.DataFrame, output_path: Path, title: str = "") -> None:
    """Writes a DataFrame to DOCX as a grid table with a bold header row."""
    doc = Document()
    if title:
        doc.add_heading(title, level=1)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        run = header_cells[i].paragraphs[0].add_run(str(column))
        run.bold = True
        run.font.size = Pt(10)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = "" if value is None else str(value)
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.save(str(output_path))


def dataframe_to_pdf(df: pd.DataFrame, output_path: Path, page_size: str = "a4", title: str = "") -> None:
    """Writes a DataFrame to PDF as a paginated table with a repeating header.

    Switches to landscape automatically when the table is wide.
    """
    base = A4 if page_size == "a4" else letter
    pagesize = landscape(base) if len(df.columns) > 6 else base

    pdf = SimpleDocTemplate(
        str(output_path), pagesize=pagesize, title=title or "Data Export",
        leftMargin=1.5 * cm, rightMargin=1.5 * cm, topMargin=1.5 * cm, bottomMargin=1.5 * cm,
    )

    def clip(value: object) -> str:
        text = "" if value is None else str(value)
        return text if len(text) <= 120 else text[:117] + "..."

    rows = [[clip(c) for c in df.columns]] + [[clip(v) for v in row] for row in df.itertuples(index=False)]

    table = Table(rows, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{_HEADER_FILL_HEX}")),
        ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor(f"#{_HEADER_TEXT_HEX}")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor(f"#{_BORDER_HEX}")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))

    pdf.build([table])
